"""Shared helpers for generating .docx files for Globalisering-momentet.

Style follows GENERERINGSPLAN.md section 7:
- A4, 1" margins
- Arial 12pt body
- Headers: H1 (18pt), H2 (14pt), H3 (12pt bold)
- Header: "Samhällskunskap 3 — Globalisering — Från vardag till världssystem" (right, grey)
- Footer: "Lektion N | Sida X" with PageNumber.CURRENT
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION


def make_document(footer_text: str = "Sida "):
    """Create a new Document with A4 + Arial defaults."""
    doc = Document()

    # A4 page size + 1" margins
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)

    # Set Normal style to Arial 12pt
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(12)
    rFonts = normal.element.rPr.rFonts
    rFonts.set(qn("w:eastAsia"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rFonts.set(qn("w:cs"), "Arial")

    # Heading styles
    for name, size, bold in [("Heading 1", 18, True), ("Heading 2", 14, True), ("Heading 3", 12, True)]:
        st = doc.styles[name]
        st.font.name = "Arial"
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        rF = st.element.rPr.rFonts
        rF.set(qn("w:eastAsia"), "Arial")
        rF.set(qn("w:hAnsi"), "Arial")
        rF.set(qn("w:cs"), "Arial")

    # Header
    header = doc.sections[0].header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_p.add_run("Samhällskunskap 3 — Globalisering — Från vardag till världssystem")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run.font.italic = True

    # Footer with page number
    footer = doc.sections[0].footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run(footer_text)
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    # Insert PAGE field
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    return doc


def add_title(doc, title, subtitle=None):
    """Add centered title + optional subtitle at top of document."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.name = "Arial"
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(subtitle)
        run2.font.name = "Arial"
        run2.font.size = Pt(11)
        run2.font.italic = True
        run2.font.color.rgb = RGBColor(0x60, 0x60, 0x70)


def add_h1(doc, text):
    h = doc.add_heading(text, level=1)
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(6)
    return h


def add_h2(doc, text):
    h = doc.add_heading(text, level=2)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    return h


def add_h3(doc, text):
    h = doc.add_heading(text, level=3)
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(2)
    return h


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.bold = bold
    run.font.italic = italic
    return p


def add_rich_para(doc, segments):
    """Add paragraph with multiple formatted segments.

    segments: list of (text, {"bold": bool, "italic": bool, "size": int}) tuples.
    """
    p = doc.add_paragraph()
    for seg in segments:
        if isinstance(seg, tuple):
            text, opts = seg
        else:
            text, opts = seg, {}
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(opts.get("size", 12))
        run.font.bold = opts.get("bold", False)
        run.font.italic = opts.get("italic", False)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.runs[0] if p.runs else p.add_run("")
    # Add text to a fresh run for consistent styling
    p.text = ""
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.text = ""
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    return p


def add_table(doc, rows, col_widths_cm=None, header_row=True):
    """Add a table. rows is a list of list of cell contents (strings).

    col_widths_cm optional list of column widths in cm.
    """
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for row in table.rows:
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    for i, row in enumerate(rows):
        for j, content in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(content)
            run.font.name = "Arial"
            run.font.size = Pt(11)
            if header_row and i == 0:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                # Shade header
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "2b2d42")
                tc_pr.append(shd)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    return table


def add_infobox(doc, title, text, color_hex="ffe66d"):
    """Add a highlighted info box (single-cell table with shading)."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.text = ""
    # Title
    title_p = cell.paragraphs[0]
    title_run = title_p.add_run(title)
    title_run.font.name = "Arial"
    title_run.font.size = Pt(11)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x30, 0x20, 0x10)
    # Body
    body_p = cell.add_paragraph()
    body_run = body_p.add_run(text)
    body_run.font.name = "Arial"
    body_run.font.size = Pt(11)
    body_run.font.color.rgb = RGBColor(0x30, 0x20, 0x10)

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)
    return table


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
